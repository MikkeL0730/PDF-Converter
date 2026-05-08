import cv2
import numpy as np
from paddleocr import PaddleOCR
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ocr = PaddleOCR(use_angle_cls=True, lang='en')  # поддерживает русский, английский, китайский и др.

def extract_font_properties(img, bbox):
    """Анализ bounding box для определения свойств шрифта"""
    x1, y1, x2, y2 = map(int, bbox[0] + bbox[2])
    roi = img[y1:y2, x1:x2]
    
    if roi.size == 0:
        return {"size_pt": 11, "bold": False, "italic": False, "color": RGBColor(0, 0, 0)}
    
    # Определение жирности через толщину штриха
    gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY)
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    if contours:
        avg_stroke = np.mean([cv2.contourArea(c) / cv2.arcLength(c, True) for c in contours if cv2.arcLength(c, True) > 0])
        is_bold = avg_stroke > 5.0  # эмпирический порог
    else:
        is_bold = False
    
    # Размер шрифта (приблизительно)
    height_pt = (y2 - y1) * 0.75  # 72 DPI -> points
    size_pt = max(8, min(72, int(height_pt * 72 / 96)))  # нормализация
    
    # Средний цвет текста
    text_mask = thresh < 128 if np.mean(thresh) > 128 else thresh > 128
    if np.any(text_mask):
        avg_color = roi[text_mask].mean(axis=0)
        color = RGBColor(int(avg_color[0]), int(avg_color[1]), int(avg_color[2]))
    else:
        color = RGBColor(0, 0, 0)
    
    return {"size_pt": size_pt, "bold": is_bold, "italic": False, "color": color}

def pdf_scanned_to_docx(pdf_path, output_path, lang='ru,en'):
    """Главная функция"""
    from pdf2image import convert_from_path
    
    images = convert_from_path(pdf_path, dpi=300, poppler_path=r'C:\Users\vanak\Downloads\Python_OCR-master\Python_OCR-master\poppler-0.68.0_x86\poppler-0.68.0\bin')
    doc = Document()
    
    for page_num, img_pil in enumerate(images):
        img_cv = cv2.cvtColor(np.array(img_pil), cv2.COLOR_RGB2BGR)
        
        # OCR распознавание
        result = ocr.ocr(img_cv, cls=True)
        
        # Группировка строк по Y координате
        lines = {}
        for line in result[0]:
            bbox, (text, confidence) = line
            if confidence < 0.7:
                continue
            
            y_center = (bbox[0][1] + bbox[2][1]) / 2
            line_key = round(y_center / 15)  # группировка строк
            
            if line_key not in lines:
                lines[line_key] = []
            lines[line_key].append((bbox, text))
        
        # Сортировка строк сверху вниз
        for line_key in sorted(lines.keys()):
            line_items = sorted(lines[line_key], key=lambda x: x[0][0][0])  # по X
            
            paragraph = doc.add_paragraph()
            
            for idx, (bbox, text) in enumerate(line_items):
                font_props = extract_font_properties(img_cv, bbox)
                
                run = paragraph.add_run(text)
                run.font.size = Pt(font_props["size_pt"])
                run.font.color.rgb = font_props["color"]
                run.bold = font_props["bold"]
                run.italic = font_props["italic"]
                
                # Добавляем пробел между словами (кроме последнего)
                if idx < len(line_items) - 1:
                    run = paragraph.add_run(" ")
                    run.font.size = Pt(font_props["size_pt"])
        
        # Разрыв страницы (кроме последней)
        if page_num < len(images) - 1:
            doc.add_page_break()
    
    doc.save(output_path)

# Использование
pdf_scanned_to_docx(r"C:\Users\vanak\Downloads\pdf_to_scan_69fdba2633707_temp_69fdba25c59d2.pdf", "output.docx", lang='ru,en')
